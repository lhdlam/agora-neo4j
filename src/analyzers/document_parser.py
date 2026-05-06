"""Document parsers for extracting text from various file formats."""

from __future__ import annotations

from abc import ABC, abstractmethod
import logging
from pathlib import Path
from typing import Any

from src.analyzers.document_models import DocumentSection, FileType, ParsedDocument

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """Base parser interface for document text extraction."""

    @abstractmethod
    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a file and return a ParsedDocument."""

    @abstractmethod
    def supports(self, file_path: Path) -> bool:
        """Check if this parser supports the given file."""


class TextParser(DocumentParser):
    """Extract text from .txt and .md files."""

    ENCODINGS = ("utf-8", "utf-8-sig", "latin-1", "cp1252")

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".txt", ".md", ".text", ".log"}

    def parse(self, file_path: Path) -> ParsedDocument:
        raw_text = self._read_with_fallback(file_path)
        sections = self._extract_sections(raw_text, file_path.suffix.lower())
        file_type: FileType = "md" if file_path.suffix.lower() == ".md" else "txt"

        return ParsedDocument(
            filename=file_path.name,
            file_type=file_type,
            raw_text=raw_text,
            metadata={"encoding": "utf-8", "size_bytes": file_path.stat().st_size},
            sections=sections,
        )

    def _read_with_fallback(self, file_path: Path) -> str:
        for encoding in self.ENCODINGS:
            try:
                return file_path.read_text(encoding=encoding)
            except (UnicodeDecodeError, ValueError):
                continue
        # Last resort: read as bytes and decode with replace
        return file_path.read_bytes().decode("utf-8", errors="replace")

    def _extract_sections(self, text: str, suffix: str) -> list[DocumentSection]:
        sections: list[DocumentSection] = []
        if suffix == ".md":
            sections = self._parse_markdown_sections(text)
        else:
            # Plain text: treat paragraphs as sections
            paragraphs = text.split("\n\n")
            for para in paragraphs:
                stripped = para.strip()
                if stripped:
                    sections.append(DocumentSection(content=stripped, level=0))
        return sections

    def _parse_markdown_sections(self, text: str) -> list[DocumentSection]:
        sections: list[DocumentSection] = []
        current_content: list[str] = []
        current_heading: str | None = None
        current_level = 0

        for line in text.split("\n"):
            if line.startswith("#"):
                # Flush previous section
                if current_content or current_heading:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n".join(current_content).strip(),
                            level=current_level,
                        )
                    )
                    current_content = []

                # Parse heading level
                level = 0
                for ch in line:
                    if ch == "#":
                        level += 1
                    else:
                        break
                current_heading = line[level:].strip()
                current_level = level
            else:
                current_content.append(line)

        # Flush last section
        if current_content or current_heading:
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content="\n".join(current_content).strip(),
                    level=current_level,
                )
            )

        return sections


class WordParser(DocumentParser):
    """Extract text from .docx files using python-docx."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".docx"}

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return ParsedDocument(
                filename=file_path.name,
                file_type="docx",
                raw_text="[Error: python-docx not installed]",
                metadata={"error": "python-docx not installed"},
            )

        doc = DocxDocument(str(file_path))
        metadata = self._extract_metadata(doc)
        sections = self._extract_sections(doc)
        raw_text = "\n\n".join(s.content for s in sections if s.content)

        return ParsedDocument(
            filename=file_path.name,
            file_type="docx",
            raw_text=raw_text,
            metadata=metadata,
            sections=sections,
        )

    def _extract_metadata(self, doc: Any) -> dict[str, Any]:
        meta: dict[str, Any] = {}
        try:
            core = doc.core_properties
            if core.author:
                meta["author"] = core.author
            if core.title:
                meta["title"] = core.title
            if core.created:
                meta["created"] = str(core.created)
            if core.modified:
                meta["modified"] = str(core.modified)
        except Exception:  # noqa: BLE001 — best effort metadata extraction
            pass
        return meta

    def _extract_sections(self, doc: Any) -> list[DocumentSection]:
        sections: list[DocumentSection] = []
        current_content: list[str] = []
        current_heading: str | None = None
        current_level = 0

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()

            if "heading" in style_name:
                # Flush previous section
                if current_content or current_heading:
                    sections.append(
                        DocumentSection(
                            heading=current_heading,
                            content="\n".join(current_content).strip(),
                            level=current_level,
                        )
                    )
                    current_content = []

                # Detect heading level from style name (e.g., "Heading 1" -> 1)
                level = 1
                for ch in style_name:
                    if ch.isdigit():
                        level = int(ch)
                        break
                current_heading = para.text.strip()
                current_level = level
            elif para.text.strip():
                current_content.append(para.text.strip())

        # Flush last section
        if current_content or current_heading:
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    content="\n".join(current_content).strip(),
                    level=current_level,
                )
            )

        # Extract tables as additional sections
        table_sections = self._extract_tables(doc)
        sections.extend(table_sections)

        return sections

    def _extract_tables(self, doc: Any) -> list[DocumentSection]:
        sections: list[DocumentSection] = []
        for i, table in enumerate(doc.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                content = "\n".join(rows)
                sections.append(
                    DocumentSection(
                        heading=f"Table {i + 1}",
                        content=content,
                        level=0,
                    )
                )
        return sections


class ImageParser(DocumentParser):
    """Extract text from images using Tesseract OCR."""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"}

    def parse(self, file_path: Path) -> ParsedDocument:
        raw_text = self._ocr_extract(file_path)

        return ParsedDocument(
            filename=file_path.name,
            file_type="image",
            raw_text=raw_text,
            metadata={
                "size_bytes": file_path.stat().st_size,
                "ocr_engine": "tesseract",
            },
            sections=[DocumentSection(content=raw_text, level=0)] if raw_text.strip() else [],
        )

    def _ocr_extract(self, file_path: Path) -> str:
        try:
            from PIL import Image
        except ImportError:
            logger.warning("Pillow not installed. Install with: pip install Pillow")
            return "[Error: Pillow not installed]"

        try:
            import pytesseract  # type: ignore[import-untyped]
        except ImportError:
            logger.warning(
                "pytesseract not installed. Install with: pip install pytesseract "
                "and install Tesseract OCR: brew install tesseract"
            )
            return "[Error: pytesseract not installed]"

        try:
            img = Image.open(file_path)
            # Use Vietnamese + English language pack if available
            try:
                text: str = pytesseract.image_to_string(img, lang="vie+eng")
            except Exception:  # noqa: BLE001 — fallback to default lang
                text = pytesseract.image_to_string(img)
            return text.strip()
        except Exception as e:  # noqa: BLE001 — graceful degradation
            logger.exception("OCR failed for %s", file_path.name)
            return f"[OCR Error: {e}]"


# --- Parser Registry ---

_PARSERS: list[DocumentParser] = [
    TextParser(),
    WordParser(),
    ImageParser(),
]


def get_parser(file_path: Path) -> DocumentParser | None:
    """Return the appropriate parser for the given file, or None if unsupported."""
    for parser in _PARSERS:
        if parser.supports(file_path):
            return parser
    return None


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".text",
    ".log",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tiff",
    ".webp",
}


def parse_document(file_path: Path) -> ParsedDocument:
    """Parse a document file and return extracted content.

    Raises ValueError if the file type is not supported.
    """
    parser = get_parser(file_path)
    if parser is None:
        msg = f"Unsupported file type: {file_path.suffix}"
        raise ValueError(msg)
    return parser.parse(file_path)
